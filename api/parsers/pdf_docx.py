#!/usr/bin/env python3
"""
PDF和Word文档解析器
"""

import logging
from pathlib import Path
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

def parse_pdf(file_path: Path) -> List[str]:
    """
    解析PDF文件
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        文本块列表
    """
    logger.info(f"开始解析PDF文件: {file_path}")
    
    try:
        import pdfplumber
        
        blocks = []
        with pdfplumber.open(str(file_path)) as pdf:
            logger.info(f"PDF页数: {len(pdf.pages)}")
            
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    lines = text.splitlines()
                    page_blocks = [line.strip() for line in lines if line.strip()]
                    blocks.extend(page_blocks)
                    logger.debug(f"第{page_num}页提取文本块: {len(page_blocks)}")
        
        logger.info(f"PDF解析完成，共提取 {len(blocks)} 个文本块")
        return blocks
        
    except ImportError:
        logger.error("pdfplumber未安装，无法解析PDF文件")
        raise Exception("PDF解析依赖未安装，请安装pdfplumber")
    except Exception as e:
        logger.error(f"PDF解析失败: {e}")
        raise

def parse_docx(file_path: Path) -> List[str]:
    """
    解析Word文档
    
    Args:
        file_path: Word文档路径
        
    Returns:
        文本块列表
    """
    logger.info(f"开始解析Word文档: {file_path}")
    
    try:
        import docx
        
        doc = docx.Document(str(file_path))
        blocks = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append(text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    blocks.append(" | ".join(row_text))
        
        logger.info(f"Word文档解析完成，共提取 {len(blocks)} 个文本块")
        return blocks
        
    except ImportError:
        logger.error("python-docx未安装，无法解析Word文档")
        raise Exception("Word解析依赖未安装，请安装python-docx")
    except Exception as e:
        logger.error(f"Word文档解析失败: {e}")
        raise

def parse_text(file_path: Path) -> List[str]:
    """
    解析纯文本文件
    
    Args:
        file_path: 文本文件路径
        
    Returns:
        文本行列表
    """
    logger.info(f"开始解析文本文件: {file_path}")
    
    try:
        # 尝试不同编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    lines = f.readlines()
                
                blocks = [line.strip() for line in lines if line.strip()]
                logger.info(f"文本文件解析完成（编码: {encoding}），共 {len(blocks)} 行")
                return blocks
                
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，使用二进制模式
        logger.warning("无法确定文件编码，使用二进制模式")
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # 尝试解码
        try:
            text = content.decode('utf-8', errors='ignore')
        except:
            text = str(content)
        
        blocks = [line.strip() for line in text.splitlines() if line.strip()]
        logger.info(f"文本文件解析完成（二进制模式），共 {len(blocks)} 行")
        return blocks
        
    except Exception as e:
        logger.error(f"文本文件解析失败: {e}")
        raise

def detect_document_structure(file_path: Path) -> Dict[str, Any]:
    """
    检测文档结构
    
    Returns:
        文档结构信息
    """
    file_ext = file_path.suffix.lower()
    
    try:
        if file_ext == '.pdf':
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                return {
                    "file_type": "pdf",
                    "pages": len(pdf.pages),
                    "has_text": any(page.extract_text() for page in pdf.pages[:3])  # 检查前3页
                }
                
        elif file_ext in ['.docx', '.doc']:
            import docx
            doc = docx.Document(str(file_path))
            return {
                "file_type": "word",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "has_content": len([p for p in doc.paragraphs if p.text.strip()]) > 0
            }
            
        elif file_ext in ['.txt', '.csv']:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            return {
                "file_type": "text",
                "lines": len(lines),
                "non_empty_lines": len([line for line in lines if line.strip()])
            }
            
        else:
            return {"file_type": "unknown", "error": f"不支持的文件类型: {file_ext}"}
            
    except Exception as e:
        logger.error(f"检测文档结构失败: {e}")
        return {"error": str(e)}

#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
测试DOCX和PDF解析功能
"""

import requests
import json
import time
from pathlib import Path
from docx import Document
from docx.shared import Inches

def create_test_docx():
    """创建测试DOCX文件"""
    print("📄 创建测试DOCX文件...")
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('测试文档标题', 0)
    
    # 添加段落
    doc.add_paragraph('这是第一个段落，包含一些测试内容。')
    doc.add_paragraph('这是第二个段落，用于验证段落解析功能。')
    
    # 添加子标题
    doc.add_heading('数据表格', level=1)
    
    # 添加表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '编号'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '描述'
    
    # 添加数据行
    for i in range(1, 4):
        row_cells = table.add_row().cells
        row_cells[0].text = f'00{i}'
        row_cells[1].text = f'项目{i}'
        row_cells[2].text = f'这是第{i}个测试项目的描述'
    
    # 添加更多段落
    doc.add_paragraph('表格后的段落内容。')
    
    # 添加列表
    doc.add_paragraph('重要事项：', style='Heading 2')
    doc.add_paragraph('• 第一个要点', style='List Bullet')
    doc.add_paragraph('• 第二个要点', style='List Bullet')
    doc.add_paragraph('• 第三个要点', style='List Bullet')
    
    # 保存文件
    file_path = Path("test_document.docx")
    doc.save(str(file_path))
    
    print(f"✅ DOCX文件创建成功: {file_path}")
    return file_path

def create_test_pdf():
    """创建测试PDF文件（使用reportlab）"""
    print("📑 创建测试PDF文件...")
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        file_path = Path("test_document.pdf")
        doc = SimpleDocTemplate(str(file_path), pagesize=letter)
        
        # 获取样式
        styles = getSampleStyleSheet()
        
        # 构建内容
        story = []
        
        # 标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
        )
        story.append(Paragraph("测试PDF文档", title_style))
        story.append(Spacer(1, 12))
        
        # 段落
        story.append(Paragraph("这是第一个段落，包含一些测试内容。", styles['Normal']))
        story.append(Paragraph("这是第二个段落，用于验证PDF解析功能。", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # 子标题
        story.append(Paragraph("数据表格", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # 表格数据
        table_data = [
            ['编号', '名称', '描述'],
            ['001', '项目1', '这是第1个测试项目的描述'],
            ['002', '项目2', '这是第2个测试项目的描述'],
            ['003', '项目3', '这是第3个测试项目的描述']
        ]
        
        # 创建表格
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 12))
        
        # 更多段落
        story.append(Paragraph("表格后的段落内容。", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # 列表
        story.append(Paragraph("重要事项：", styles['Heading3']))
        story.append(Paragraph("• 第一个要点", styles['Normal']))
        story.append(Paragraph("• 第二个要点", styles['Normal']))
        story.append(Paragraph("• 第三个要点", styles['Normal']))
        
        # 生成PDF
        doc.build(story)
        
        print(f"✅ PDF文件创建成功: {file_path}")
        return file_path
        
    except ImportError:
        print("❌ reportlab未安装，跳过PDF创建")
        return None

def test_document_parsing(file_path: Path, file_type: str):
    """测试文档解析"""
    print(f"\n🧪 测试{file_type.upper()}解析...")
    
    base_url = "http://127.0.0.1:8000"
    
    try:
        # 1. 上传文件
        print("1️⃣ 上传文件...")
        mime_types = {
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'pdf': 'application/pdf'
        }
        
        with open(file_path, 'rb') as f:
            files = {'file': (file_path.name, f, mime_types.get(file_type, 'application/octet-stream'))}
            response = requests.post(f"{base_url}/kg/upload", files=files)
        
        if response.status_code != 200:
            print(f"❌ 上传失败: {response.status_code}")
            return False
            
        upload_data = response.json()
        upload_id = upload_data['upload_id']
        print(f"✅ 上传成功 (ID: {upload_id})")
        
        # 2. 触发解析
        print("2️⃣ 触发解析...")
        parse_response = requests.post(f"{base_url}/kg/files/{upload_id}/parse")
        
        if parse_response.status_code != 200:
            print(f"❌ 解析触发失败: {parse_response.status_code}")
            print(f"响应: {parse_response.text}")
            return False
        
        print("✅ 解析任务已启动")
        
        # 3. 等待解析完成
        print("3️⃣ 等待解析完成...")
        max_wait = 30
        wait_time = 0
        
        while wait_time < max_wait:
            status_response = requests.get(f"{base_url}/kg/files/{upload_id}/status")
            if status_response.status_code == 200:
                status_data = status_response.json()
                current_status = status_data.get('data', {}).get('status', 'unknown')
                print(f"   [{wait_time}s] 状态: {current_status}")
                
                if current_status == 'parsed':
                    print("✅ 解析完成")
                    break
                elif current_status == 'failed':
                    error_msg = status_data.get('data', {}).get('error', '未知错误')
                    print(f"❌ 解析失败: {error_msg}")
                    return False
            
            time.sleep(2)
            wait_time += 2
        
        if wait_time >= max_wait:
            print("❌ 解析超时")
            return False
        
        # 4. 获取解析结果
        print("4️⃣ 获取解析结果...")
        preview_response = requests.get(f"{base_url}/kg/files/{upload_id}/preview")
        
        if preview_response.status_code != 200:
            print(f"❌ 获取结果失败: {preview_response.status_code}")
            return False
            
        preview_data = preview_response.json()
        
        if not preview_data.get('success'):
            print(f"❌ 解析结果无效: {preview_data.get('error', '未知错误')}")
            return False
        
        # 5. 分析解析结果
        print("5️⃣ 分析解析结果...")
        raw_data = preview_data.get('data', {}).get('raw_data', [])
        metadata = preview_data.get('data', {}).get('metadata', {})
        
        print(f"   总记录数: {len(raw_data)}")
        print(f"   元数据字段: {len(metadata)}")
        
        if len(raw_data) == 0:
            print("❌ 解析结果为空！")
            return False
        
        # 显示前几条记录
        print(f"   前3条记录:")
        for i, record in enumerate(raw_data[:3]):
            print(f"      记录{i+1}: {record.get('content_type', 'unknown')} - {record.get('content', '')[:50]}...")
        
        # 统计内容类型
        content_types = {}
        for record in raw_data:
            content_type = record.get('content_type', 'unknown')
            content_types[content_type] = content_types.get(content_type, 0) + 1
        
        print(f"   内容类型统计: {content_types}")
        
        return True
        
    except Exception as e:
        print(f"❌ 测试异常: {e}")
        return False

def main():
    """主测试函数"""
    print("🔍 开始DOCX和PDF解析测试")
    print("=" * 60)
    
    # 测试DOCX
    docx_file = create_test_docx()
    docx_success = test_document_parsing(docx_file, 'docx')
    
    # 测试PDF
    pdf_file = create_test_pdf()
    if pdf_file:
        pdf_success = test_document_parsing(pdf_file, 'pdf')
    else:
        pdf_success = False
        print("⚠️ PDF测试跳过（缺少依赖）")
    
    # 清理测试文件
    print(f"\n🧹 清理测试文件...")
    if docx_file.exists():
        docx_file.unlink()
        print(f"   删除: {docx_file.name}")
    
    if pdf_file and pdf_file.exists():
        pdf_file.unlink()
        print(f"   删除: {pdf_file.name}")
    
    # 总结
    print(f"\n📊 测试结果总结:")
    print(f"   DOCX解析: {'✅ 成功' if docx_success else '❌ 失败'}")
    print(f"   PDF解析: {'✅ 成功' if pdf_success else '❌ 失败'}")
    
    if docx_success and pdf_success:
        print("\n🎉 所有测试通过！")
    elif docx_success or pdf_success:
        print("\n⚠️ 部分测试通过，需要进一步调试")
    else:
        print("\n❌ 所有测试失败，需要检查解析器实现")

if __name__ == "__main__":
    main()

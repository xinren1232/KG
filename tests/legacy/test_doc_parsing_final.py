#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文档解析功能
"""

import requests
import json
import time
import pandas as pd
import docx
from pathlib import Path

def test_excel_parsing():
    """测试Excel文档解析"""
    print("=== 测试Excel文档解析 ===")
    
    # 创建测试Excel文件
    data = {
        '工厂名称': ['索尼', '苹果', '中兴'],
        '产品型号': ['XM4', 'iPhone14', 'Axon30'],
        '发现时间': pd.to_datetime(['2025-01-15 10:30:00', '2025-01-16 14:20:00', '2025-01-17 09:15:00']),
        '状态': ['已解决', '处理中', '待分析'],
        '问题描述': ['电池续航异常', '屏幕显示故障', '摄像头模糊']
    }
    df = pd.DataFrame(data)
    excel_file = 'test_excel_parsing.xlsx'
    df.to_excel(excel_file, index=False)
    print(f"✅ 创建Excel测试文件: {excel_file}")
    
    # 上传并解析
    with open(excel_file, 'rb') as f:
        files = {'file': f}
        upload_r = requests.post('http://localhost:8000/kg/upload', files=files)
    
    print(f"上传状态: {upload_r.status_code}")
    if upload_r.status_code == 200:
        upload_result = upload_r.json()
        if upload_result.get('success'):
            upload_id = upload_result.get('upload_id')
            print(f"✅ 上传成功，ID: {upload_id}")
            
            # 触发解析
            parse_r = requests.post(f'http://localhost:8000/kg/files/{upload_id}/parse')
            if parse_r.status_code == 200:
                parse_result = parse_r.json()
                if parse_result.get('success'):
                    print("✅ 解析触发成功")
                    
                    # 等待解析完成
                    time.sleep(3)
                    
                    # 获取解析结果
                    preview_r = requests.get(f'http://localhost:8000/kg/files/{upload_id}/preview')
                    if preview_r.status_code == 200:
                        preview_result = preview_r.json()
                        if preview_result.get('success'):
                            data = preview_result.get('data', {})
                            raw_data = data.get('raw_data', [])
                            entities = data.get('entities', [])
                            metadata = data.get('metadata', {})
                            
                            print(f"✅ Excel解析成功!")
                            print(f"   原始数据: {len(raw_data)} 条")
                            print(f"   识别实体: {len(entities)} 个")
                            print(f"   文件类型: {metadata.get('file_type')}")
                            
                            if raw_data:
                                print("   数据示例:")
                                for i, item in enumerate(raw_data[:2]):
                                    print(f"      {i+1}. {item.get('content', '')[:50]}...")
                            
                            return True
                        else:
                            print(f"❌ 预览失败: {preview_result.get('message')}")
                    else:
                        print(f"❌ 预览请求失败: {preview_r.text}")
                else:
                    print(f"❌ 解析失败: {parse_result.get('message')}")
            else:
                print(f"❌ 解析请求失败: {parse_r.text}")
        else:
            print(f"❌ 上传失败: {upload_result.get('message')}")
    else:
        print(f"❌ 上传请求失败: {upload_r.text}")
    
    return False

def test_word_parsing():
    """测试Word文档解析"""
    print("\n=== 测试Word文档解析 ===")
    
    # 创建测试Word文档
    doc = docx.Document()
    doc.add_heading('硬件质量测试报告', 0)
    doc.add_paragraph('本报告描述了智能手机硬件的质量测试结果。')
    doc.add_paragraph('电池续航测试发现异常，续航时间不足预期。')
    doc.add_paragraph('屏幕显示测试正常，无故障现象。')
    doc.add_paragraph('摄像头功能测试发现部分功能异常。')
    
    # 添加表格
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '测试项目'
    hdr_cells[1].text = '测试结果'
    hdr_cells[2].text = '备注'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '电池续航'
    row_cells[1].text = '异常'
    row_cells[2].text = '续航时间不足'
    
    word_file = 'test_word_parsing.docx'
    doc.save(word_file)
    print(f"✅ 创建Word测试文件: {word_file}")
    
    # 上传并解析
    with open(word_file, 'rb') as f:
        files = {'file': f}
        upload_r = requests.post('http://localhost:8000/kg/upload', files=files)
    
    print(f"上传状态: {upload_r.status_code}")
    if upload_r.status_code == 200:
        upload_result = upload_r.json()
        if upload_result.get('success'):
            upload_id = upload_result.get('upload_id')
            print(f"✅ 上传成功，ID: {upload_id}")
            
            # 触发解析
            parse_r = requests.post(f'http://localhost:8000/kg/files/{upload_id}/parse')
            if parse_r.status_code == 200:
                parse_result = parse_r.json()
                if parse_result.get('success'):
                    print("✅ 解析触发成功")
                    
                    # 等待解析完成
                    time.sleep(3)
                    
                    # 获取解析结果
                    preview_r = requests.get(f'http://localhost:8000/kg/files/{upload_id}/preview')
                    if preview_r.status_code == 200:
                        preview_result = preview_r.json()
                        if preview_result.get('success'):
                            data = preview_result.get('data', {})
                            raw_data = data.get('raw_data', [])
                            entities = data.get('entities', [])
                            metadata = data.get('metadata', {})
                            
                            print(f"✅ Word解析成功!")
                            print(f"   原始数据: {len(raw_data)} 条")
                            print(f"   识别实体: {len(entities)} 个")
                            print(f"   文件类型: {metadata.get('file_type')}")
                            
                            if raw_data:
                                print("   内容示例:")
                                for i, item in enumerate(raw_data[:3]):
                                    content = item.get('content', '')[:60]
                                    item_type = item.get('type', '未知')
                                    print(f"      {i+1}. [{item_type}] {content}...")
                            
                            if entities:
                                print("   识别实体:")
                                for entity in entities[:3]:
                                    name = entity.get('name')
                                    entity_type = entity.get('type')
                                    confidence = entity.get('confidence', 0)
                                    print(f"      - {name} ({entity_type}) - 置信度: {confidence}")
                            
                            return True
                        else:
                            print(f"❌ 预览失败: {preview_result.get('message')}")
                    else:
                        print(f"❌ 预览请求失败: {preview_r.text}")
                else:
                    print(f"❌ 解析失败: {parse_result.get('message')}")
            else:
                print(f"❌ 解析请求失败: {parse_r.text}")
        else:
            print(f"❌ 上传失败: {upload_result.get('message')}")
    else:
        print(f"❌ 上传请求失败: {upload_r.text}")
    
    return False

def test_text_parsing():
    """测试文本文档解析"""
    print("\n=== 测试文本文档解析 ===")
    
    # 创建测试文本文件
    text_content = """硬件质量测试报告

测试概述：
本次测试针对智能手机的主要硬件组件进行了全面的质量检测。

测试项目：
1. 电池续航测试 - 发现异常
2. 屏幕显示测试 - 正常
3. 摄像头功能测试 - 部分异常
4. 充电接口测试 - 正常
5. 扬声器测试 - 正常

问题分析：
电池模块存在续航时间不足的问题，需要进一步检测。
摄像头在低光环境下存在对焦异常。

建议措施：
1. 更换电池供应商
2. 优化摄像头算法
3. 加强质量控制流程
"""
    
    text_file = 'test_text_parsing.txt'
    with open(text_file, 'w', encoding='utf-8') as f:
        f.write(text_content)
    print(f"✅ 创建文本测试文件: {text_file}")
    
    # 上传并解析
    with open(text_file, 'rb') as f:
        files = {'file': f}
        upload_r = requests.post('http://localhost:8000/kg/upload', files=files)
    
    print(f"上传状态: {upload_r.status_code}")
    if upload_r.status_code == 200:
        upload_result = upload_r.json()
        if upload_result.get('success'):
            upload_id = upload_result.get('upload_id')
            print(f"✅ 上传成功，ID: {upload_id}")
            
            # 触发解析
            parse_r = requests.post(f'http://localhost:8000/kg/files/{upload_id}/parse')
            if parse_r.status_code == 200:
                parse_result = parse_r.json()
                if parse_result.get('success'):
                    print("✅ 解析触发成功")
                    
                    # 等待解析完成
                    time.sleep(2)
                    
                    # 获取解析结果
                    preview_r = requests.get(f'http://localhost:8000/kg/files/{upload_id}/preview')
                    if preview_r.status_code == 200:
                        preview_result = preview_r.json()
                        if preview_result.get('success'):
                            data = preview_result.get('data', {})
                            raw_data = data.get('raw_data', [])
                            entities = data.get('entities', [])
                            metadata = data.get('metadata', {})
                            
                            print(f"✅ 文本解析成功!")
                            print(f"   原始数据: {len(raw_data)} 条")
                            print(f"   识别实体: {len(entities)} 个")
                            print(f"   文件类型: {metadata.get('file_type')}")
                            
                            if entities:
                                print("   识别实体:")
                                for entity in entities[:5]:
                                    name = entity.get('name')
                                    entity_type = entity.get('type')
                                    confidence = entity.get('confidence', 0)
                                    print(f"      - {name} ({entity_type}) - 置信度: {confidence}")
                            
                            return True
                        else:
                            print(f"❌ 预览失败: {preview_result.get('message')}")
                    else:
                        print(f"❌ 预览请求失败: {preview_r.text}")
                else:
                    print(f"❌ 解析失败: {parse_result.get('message')}")
            else:
                print(f"❌ 解析请求失败: {parse_r.text}")
        else:
            print(f"❌ 上传失败: {upload_result.get('message')}")
    else:
        print(f"❌ 上传请求失败: {upload_r.text}")
    
    return False

def main():
    """主测试函数"""
    print("🔍 开始测试文档解析功能...")
    
    # 测试API连接
    try:
        r = requests.get('http://localhost:8000/health', timeout=5)
        if r.status_code == 200:
            print("✅ API服务器连接正常")
        else:
            print(f"❌ API服务器响应异常: {r.status_code}")
            return
    except Exception as e:
        print(f"❌ 无法连接到API服务器: {e}")
        return
    
    # 运行各种文档类型的测试
    results = []
    
    # 测试Excel解析
    results.append(("Excel", test_excel_parsing()))
    
    # 测试Word解析
    results.append(("Word", test_word_parsing()))
    
    # 测试文本解析
    results.append(("Text", test_text_parsing()))
    
    # 总结测试结果
    print("\n" + "="*50)
    print("📊 测试结果总结:")
    print("="*50)
    
    success_count = 0
    for doc_type, success in results:
        status = "✅ 成功" if success else "❌ 失败"
        print(f"   {doc_type}文档解析: {status}")
        if success:
            success_count += 1
    
    print(f"\n总体成功率: {success_count}/{len(results)} ({success_count/len(results)*100:.1f}%)")
    
    if success_count == len(results):
        print("🎉 所有文档类型解析功能正常!")
    else:
        print("⚠️  部分文档类型解析存在问题，需要进一步排查。")

if __name__ == "__main__":
    main()

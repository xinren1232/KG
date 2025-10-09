#!/usr/bin/env python3
import requests
import time

def test_simple_document():
    """测试简单文档解析"""
    
    print("=== 测试文档解析修复 ===")
    
    # 创建简单的Word文档
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('测试报告', 0)
        doc.add_paragraph('这是一个测试段落，包含硬件质量相关内容。')
        doc.add_paragraph('电池续航测试结果显示存在问题。')
        doc.add_paragraph('屏幕显示正常，无异常发现。')
        
        # 添加表格
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '测试项目'
        hdr_cells[1].text = '结果'
        
        row_cells = table.add_row().cells
        row_cells[0].text = '电池测试'
        row_cells[1].text = '失败'
        
        doc.save('simple_test.docx')
        print("✅ 创建Word测试文档")
        
        # 测试上传和解析
        print("\n📄 测试Word文档解析...")
        
        # 1. 上传文件
        with open('simple_test.docx', 'rb') as f:
            files = {'file': f}
            upload_response = requests.post('http://localhost:8000/kg/upload', files=files, timeout=10)
        
        print(f"上传响应状态: {upload_response.status_code}")
        if upload_response.status_code == 200:
            upload_result = upload_response.json()
            print(f"上传结果: {upload_result}")
            
            if upload_result.get('success'):
                upload_id = upload_result.get('upload_id')
                print(f"✅ 文件上传成功: {upload_id}")
                
                # 2. 触发解析
                parse_response = requests.post(f'http://localhost:8000/kg/files/{upload_id}/parse', timeout=10)
                print(f"解析响应状态: {parse_response.status_code}")
                
                if parse_response.status_code == 200:
                    parse_result = parse_response.json()
                    print(f"解析结果: {parse_result}")
                    
                    if parse_result.get('success'):
                        print("✅ 解析触发成功")
                        
                        # 3. 等待并获取结果
                        time.sleep(3)
                        preview_response = requests.get(f'http://localhost:8000/kg/files/{upload_id}/preview', timeout=10)
                        
                        if preview_response.status_code == 200:
                            preview_result = preview_response.json()
                            print(f"预览结果: {preview_result}")
                            
                            if preview_result.get('success'):
                                data = preview_result.get('data', {})
                                raw_data = data.get('raw_data', [])
                                entities = data.get('entities', [])
                                
                                print(f"✅ 解析成功!")
                                print(f"   原始数据条数: {len(raw_data)}")
                                print(f"   识别实体数量: {len(entities)}")
                                
                                if raw_data:
                                    print("   内容示例:")
                                    for i, item in enumerate(raw_data[:3]):
                                        content = item.get('content', '')[:80]
                                        item_type = item.get('type', '未知')
                                        print(f"      {i+1}. [{item_type}] {content}...")
                                
                                if entities:
                                    print("   识别的实体:")
                                    for entity in entities[:3]:
                                        name = entity.get('name')
                                        entity_type = entity.get('type')
                                        print(f"      - {name} ({entity_type})")
                                
                                return True
                            else:
                                print(f"❌ 预览失败: {preview_result.get('message')}")
                        else:
                            print(f"❌ 预览请求失败: {preview_response.status_code}")
                    else:
                        print(f"❌ 解析失败: {parse_result.get('message')}")
                else:
                    print(f"❌ 解析请求失败: {parse_response.status_code}")
                    print(f"错误: {parse_response.text}")
            else:
                print(f"❌ 上传失败: {upload_result.get('message')}")
        else:
            print(f"❌ 上传请求失败: {upload_response.status_code}")
            print(f"错误: {upload_response.text}")
            
    except ImportError:
        print("❌ python-docx未安装")
    except Exception as e:
        print(f"❌ 测试异常: {e}")
        import traceback
        traceback.print_exc()
    
    return False

if __name__ == "__main__":
    test_simple_document()

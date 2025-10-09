#!/usr/bin/env python3
import requests
import time
from pathlib import Path

def create_test_documents():
    """创建测试文档"""
    
    # 创建测试Word文档
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('硬件质量测试报告', 0)
        
        doc.add_heading('1. 测试概述', level=1)
        doc.add_paragraph('本报告描述了对索尼XM4耳机的质量测试结果。')
        
        doc.add_heading('2. 测试项目', level=1)
        doc.add_paragraph('主要测试项目包括：')
        doc.add_paragraph('• 电池续航测试')
        doc.add_paragraph('• 音质测试')
        doc.add_paragraph('• 降噪效果测试')
        doc.add_paragraph('• 连接稳定性测试')
        
        doc.add_heading('3. 测试结果', level=1)
        
        # 添加表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '测试项目'
        hdr_cells[1].text = '结果'
        hdr_cells[2].text = '备注'
        
        # 添加数据行
        test_data = [
            ('电池续航', '通过', '续航时间达到30小时'),
            ('音质测试', '通过', '频响范围符合标准'),
            ('降噪效果', '部分通过', '在高频段有轻微问题'),
            ('连接稳定性', '失败', '在2.4GHz频段存在干扰')
        ]
        
        for item, result, note in test_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = result
            row_cells[2].text = note
        
        doc.add_heading('4. 问题分析', level=1)
        doc.add_paragraph('发现的主要问题：')
        doc.add_paragraph('1. 连接稳定性问题可能由于天线设计缺陷导致')
        doc.add_paragraph('2. 降噪算法在高频段需要优化')
        
        doc.save('test_hardware_report.docx')
        print("✅ 创建Word测试文档: test_hardware_report.docx")
        
    except ImportError:
        print("❌ python-docx未安装，无法创建Word测试文档")
    
    # 创建测试PDF文档（使用reportlab）
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate("test_quality_analysis.pdf", pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # 标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
        )
        story.append(Paragraph("质量分析报告", title_style))
        story.append(Spacer(1, 12))
        
        # 内容
        story.append(Paragraph("1. 产品概述", styles['Heading2']))
        story.append(Paragraph("本报告分析了苹果iPhone14的质量问题。", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("2. 问题统计", styles['Heading2']))
        
        # 表格数据
        data = [
            ['问题类型', '数量', '严重程度', '状态'],
            ['屏幕显示异常', '15', '高', '处理中'],
            ['电池续航短', '8', '中', '已解决'],
            ['摄像头模糊', '12', '中', '待分析'],
            ['充电接口松动', '5', '低', '已解决']
        ]
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("3. 分析结论", styles['Heading2']))
        story.append(Paragraph("主要质量问题集中在显示和摄像头模块，需要重点关注供应商质量控制。", styles['Normal']))
        
        doc.build(story)
        print("✅ 创建PDF测试文档: test_quality_analysis.pdf")
        
    except ImportError:
        print("❌ reportlab未安装，无法创建PDF测试文档")
        # 创建简单的文本文件作为替代
        with open('test_quality_analysis.txt', 'w', encoding='utf-8') as f:
            f.write("""质量分析报告

1. 产品概述
本报告分析了苹果iPhone14的质量问题。

2. 问题统计
屏幕显示异常: 15个问题，严重程度高，处理中
电池续航短: 8个问题，严重程度中，已解决
摄像头模糊: 12个问题，严重程度中，待分析
充电接口松动: 5个问题，严重程度低，已解决

3. 分析结论
主要质量问题集中在显示和摄像头模块，需要重点关注供应商质量控制。
""")
        print("✅ 创建文本测试文档: test_quality_analysis.txt")

def test_document_parsing():
    """测试文档解析功能"""
    
    print("=== 测试文档解析修复 ===")
    
    # 创建测试文档
    create_test_documents()
    
    # 测试Word文档解析
    if Path('test_hardware_report.docx').exists():
        print("\n📄 测试Word文档解析...")
        test_file_parsing('test_hardware_report.docx', 'Word')
    
    # 测试PDF文档解析
    if Path('test_quality_analysis.pdf').exists():
        print("\n📄 测试PDF文档解析...")
        test_file_parsing('test_quality_analysis.pdf', 'PDF')
    elif Path('test_quality_analysis.txt').exists():
        print("\n📄 测试文本文档解析...")
        test_file_parsing('test_quality_analysis.txt', '文本')

def test_file_parsing(filename, file_type):
    """测试单个文件的解析"""
    try:
        # 1. 上传文件
        with open(filename, 'rb') as f:
            files = {'file': f}
            upload_response = requests.post('http://localhost:8000/kg/upload', files=files)
        
        if upload_response.status_code != 200:
            print(f"   ❌ {file_type}文件上传失败: {upload_response.status_code}")
            return
        
        upload_result = upload_response.json()
        if not upload_result.get('success'):
            print(f"   ❌ {file_type}文件上传失败: {upload_result.get('message')}")
            return
        
        upload_id = upload_result.get('upload_id')
        print(f"   ✅ {file_type}文件上传成功: {upload_id}")
        
        # 2. 触发解析
        parse_response = requests.post(f'http://localhost:8000/kg/files/{upload_id}/parse')
        
        if parse_response.status_code != 200:
            print(f"   ❌ {file_type}解析失败: {parse_response.status_code}")
            print(f"   错误: {parse_response.text}")
            return
        
        parse_result = parse_response.json()
        if not parse_result.get('success'):
            print(f"   ❌ {file_type}解析失败: {parse_result.get('message')}")
            return
        
        print(f"   ✅ {file_type}解析成功")
        
        # 3. 等待解析完成并获取结果
        time.sleep(3)
        
        preview_response = requests.get(f'http://localhost:8000/kg/files/{upload_id}/preview')
        
        if preview_response.status_code != 200:
            print(f"   ❌ 获取{file_type}解析结果失败: {preview_response.status_code}")
            return
        
        preview_result = preview_response.json()
        if not preview_result.get('success'):
            print(f"   ❌ 获取{file_type}解析结果失败: {preview_result.get('message')}")
            return
        
        data = preview_result.get('data', {})
        raw_data = data.get('raw_data', [])
        entities = data.get('entities', [])
        
        print(f"   ✅ {file_type}解析结果获取成功")
        print(f"   📊 解析统计:")
        print(f"      - 原始数据条数: {len(raw_data)}")
        print(f"      - 识别实体数量: {len(entities)}")
        
        # 显示前几条解析数据
        if raw_data:
            print(f"   📄 {file_type}内容示例:")
            for i, item in enumerate(raw_data[:3]):
                content = item.get('content', '')[:100]
                item_type = item.get('type', '未知')
                print(f"      {i+1}. [{item_type}] {content}...")
        
        # 显示识别的实体
        if entities:
            print(f"   🏷️ 识别的实体:")
            for entity in entities[:5]:
                name = entity.get('name')
                entity_type = entity.get('type')
                confidence = entity.get('confidence', 0)
                print(f"      - {name} ({entity_type}) - 置信度: {confidence:.2f}")
        
        return True
        
    except Exception as e:
        print(f"   ❌ {file_type}文档解析测试异常: {e}")
        return False

if __name__ == "__main__":
    test_document_parsing()

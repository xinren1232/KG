#!/usr/bin/env python3
import requests
import time
import docx

def create_test_word():
    """创建测试Word文档"""
    doc = docx.Document()
    doc.add_heading('硬件质量测试报告', 0)
    doc.add_paragraph('本报告描述了对手机电池的质量测试。')
    doc.add_paragraph('测试发现电池续航存在异常，需要进一步分析。')
    doc.add_paragraph('屏幕显示功能正常，无故障现象。')
    
    # 添加表格
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '测试项目'
    hdr_cells[1].text = '结果'
    hdr_cells[2].text = '备注'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '电池测试'
    row_cells[1].text = '失败'
    row_cells[2].text = '续航时间不足'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '屏幕测试'
    row_cells[1].text = '通过'
    row_cells[2].text = '显示正常'
    
    doc.save('test_report.docx')
    print("✅ 创建Word测试文档: test_report.docx")

def test_word_parsing():
    """测试Word文档解析"""
    print("\n=== 测试Word文档解析 ===")
    
    # 创建测试文档
    create_test_word()
    
    try:
        # 1. 上传文件
        print("1. 上传Word文档...")
        with open('test_report.docx', 'rb') as f:
            files = {'file': f}
            upload_response = requests.post('http://localhost:8000/kg/upload', files=files, timeout=10)
        
        if upload_response.status_code != 200:
            print(f"   ❌ 上传失败: {upload_response.status_code}")
            print(f"   错误: {upload_response.text}")
            return False
        
        upload_result = upload_response.json()
        if not upload_result.get('success'):
            print(f"   ❌ 上传失败: {upload_result.get('message')}")
            return False
        
        upload_id = upload_result.get('upload_id')
        print(f"   ✅ 上传成功: {upload_id}")
        
        # 2. 触发解析
        print("2. 触发解析...")
        parse_response = requests.post(f'http://localhost:8000/kg/files/{upload_id}/parse', timeout=10)
        
        if parse_response.status_code != 200:
            print(f"   ❌ 解析请求失败: {parse_response.status_code}")
            print(f"   错误: {parse_response.text}")
            return False
        
        parse_result = parse_response.json()
        if not parse_result.get('success'):
            print(f"   ❌ 解析失败: {parse_result.get('message')}")
            return False
        
        print(f"   ✅ 解析触发成功: {parse_result.get('message')}")
        
        # 3. 等待解析完成
        print("3. 等待解析完成...")
        time.sleep(3)
        
        # 4. 获取解析结果
        print("4. 获取解析结果...")
        preview_response = requests.get(f'http://localhost:8000/kg/files/{upload_id}/preview', timeout=10)
        
        if preview_response.status_code != 200:
            print(f"   ❌ 获取结果失败: {preview_response.status_code}")
            print(f"   错误: {preview_response.text}")
            return False
        
        preview_result = preview_response.json()
        if not preview_result.get('success'):
            print(f"   ❌ 获取结果失败: {preview_result.get('message')}")
            return False
        
        data = preview_result.get('data', {})
        raw_data = data.get('raw_data', [])
        entities = data.get('entities', [])
        
        print(f"   ✅ 解析结果获取成功!")
        print(f"   📊 统计信息:")
        print(f"      - 原始数据条数: {len(raw_data)}")
        print(f"      - 识别实体数量: {len(entities)}")
        
        # 显示解析内容
        if raw_data:
            print(f"   📄 解析内容示例:")
            for i, item in enumerate(raw_data[:5]):
                content = item.get('content', '')[:80]
                item_type = item.get('type', '未知')
                print(f"      {i+1}. [{item_type}] {content}...")
        
        # 显示识别的实体
        if entities:
            print(f"   🏷️ 识别的实体:")
            for entity in entities[:5]:
                name = entity.get('name')
                entity_type = entity.get('type')
                confidence = entity.get('confidence', 0)
                print(f"      - {name} ({entity_type}) - 置信度: {confidence:.2f}")
        
        return len(raw_data) > 0  # 如果有解析数据就算成功
        
    except Exception as e:
        print(f"   ❌ 测试异常: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_api_connection():
    """测试API连接"""
    print("=== 测试API连接 ===")
    try:
        response = requests.get('http://localhost:8000/health', timeout=5)
        if response.status_code == 200:
            print("✅ API连接正常")
            return True
        else:
            print(f"❌ API响应异常: {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ API连接失败: {e}")
        return False

if __name__ == "__main__":
    print("🔧 开始测试文档解析修复...")
    
    # 测试API连接
    if not test_api_connection():
        print("❌ API连接失败，请确保API服务正在运行")
        exit(1)
    
    # 测试Word文档解析
    if test_word_parsing():
        print("\n🎉 Word文档解析修复成功!")
    else:
        print("\n❌ Word文档解析仍有问题")
    
    print("\n测试完成。")

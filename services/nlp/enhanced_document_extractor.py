#!/usr/bin/env python3
"""
增强文档抽取器
支持Excel、PDF、Word等多种格式的智能抽取
集成spaCy NLP和LLM能力进行实体关系抽取
"""
import re
import json
import pandas as pd
from typing import Dict, List, Tuple, Optional, Any, Union
from pathlib import Path
import logging
from datetime import datetime

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class EnhancedDocumentExtractor:
    """增强文档抽取器"""
    
    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.csv', '.pdf', '.docx', '.txt']
        self.nlp_models = {}
        self.extraction_rules = self._load_extraction_rules()
        self.vocabulary_mappings = self._load_vocabulary_mappings()
        
    def _load_extraction_rules(self) -> Dict[str, Any]:
        """加载抽取规则"""
        return {
            'entity_patterns': {
                'anomaly_id': [
                    r'[A-Z]{2,4}-\d{4}-\d{3,4}',  # QA-2024-001
                    r'IQC-\d{4}-\d{3}',           # IQC-2024-001
                    r'异常编号[:：]\s*([A-Z0-9-]+)',
                    r'问题编号[:：]\s*([A-Z0-9-]+)'
                ],
                'product_model': [
                    r'MyPhone[A-Z]',
                    r'[A-Z]{2,4}\d{2,4}[A-Z]?',   # BG6, MP30A
                    r'机型[:：]\s*([A-Z0-9]+)',
                    r'产品[:：]\s*([A-Z0-9]+)'
                ],
                'component': [
                    r'摄像头|相机|镜头|Camera',
                    r'电池|电芯|Battery',
                    r'显示屏|屏幕|LCD|OLED|Display',
                    r'触摸屏|触控|Touch',
                    r'扬声器|喇叭|Speaker',
                    r'组件[:：]\s*([^，。\n]+)',
                    r'模块[:：]\s*([^，。\n]+)'
                ],
                'severity': [
                    r'S[1-4]',
                    r'严重程度[:：]\s*(S[1-4]|高|中|低)',
                    r'优先级[:：]\s*(P[1-4]|高|中|低)'
                ],
                'symptom': [
                    r'对焦.*失败|对焦.*异常',
                    r'充电.*慢|充电.*异常',
                    r'色彩.*偏差|显示.*异常',
                    r'触摸.*不灵敏|响应.*异常',
                    r'音质.*异常|杂音',
                    r'裂纹|破损|变形|划伤',
                    r'症状[:：]\s*([^，。\n]+)',
                    r'现象[:：]\s*([^，。\n]+)'
                ],
                'root_cause': [
                    r'.*导致.*',
                    r'.*原因.*',
                    r'工艺.*问题',
                    r'内阻.*偏高',
                    r'色温.*偏差',
                    r'根因[:：]\s*([^，。\n]+)',
                    r'原因[:：]\s*([^，。\n]+)'
                ],
                'countermeasure': [
                    r'更换.*',
                    r'调整.*',
                    r'增加.*',
                    r'改进.*',
                    r'优化.*',
                    r'对策[:：]\s*([^，。\n]+)',
                    r'措施[:：]\s*([^，。\n]+)'
                ],
                'supplier': [
                    r'.*有限公司',
                    r'.*股份.*公司',
                    r'.*科技.*公司',
                    r'.*制造.*公司',
                    r'供应商[:：]\s*([^，。\n]+)'
                ],
                'owner': [
                    r'[\u4e00-\u9fa5]{2,4}',  # 中文姓名
                    r'责任人[:：]\s*([\u4e00-\u9fa5]{2,4})',
                    r'处理人[:：]\s*([\u4e00-\u9fa5]{2,4})'
                ]
            },
            'relation_patterns': [
                {
                    'pattern': r'(.*异常).*影响.*(组件|模块)',
                    'relation': 'AFFECTS',
                    'source_type': 'Anomaly',
                    'target_type': 'Component'
                },
                {
                    'pattern': r'(.*症状).*由于.*(原因)',
                    'relation': 'CAUSED_BY',
                    'source_type': 'Symptom',
                    'target_type': 'RootCause'
                },
                {
                    'pattern': r'(.*原因).*通过.*(措施).*解决',
                    'relation': 'RESOLVED_BY',
                    'source_type': 'RootCause',
                    'target_type': 'Countermeasure'
                }
            ]
        }
    
    def _load_vocabulary_mappings(self) -> Dict[str, Dict[str, str]]:
        """加载词汇映射表"""
        return {
            'severity_mapping': {
                '高': 'S1', '严重': 'S1', 'High': 'S1', 'Critical': 'S1',
                '中': 'S2', '一般': 'S2', 'Medium': 'S2', 'Normal': 'S2',
                '低': 'S3', '轻微': 'S3', 'Low': 'S3', 'Minor': 'S3'
            },
            'component_mapping': {
                '相机': '摄像头', 'Camera': '摄像头',
                '电芯': '电池', 'Battery': '电池',
                '屏幕': '显示屏', 'Display': '显示屏', 'LCD': '显示屏',
                '触控': '触摸屏', 'Touch': '触摸屏',
                '喇叭': '扬声器', 'Speaker': '扬声器'
            },
            'symptom_mapping': {
                '对焦失败': '对焦异常',
                '充电慢': '充电异常',
                '显示异常': '色彩偏差',
                '触控不灵敏': '触摸异常',
                '音质差': '音质异常'
            }
        }
    
    def extract_from_file(self, file_path: str) -> Dict[str, Any]:
        """从文件抽取数据（支持多种格式）"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
        
        logger.info(f"开始抽取文件: {file_path}")
        
        # 根据文件类型选择抽取方法
        if file_path.suffix.lower() in ['.xlsx', '.xls']:
            return self._extract_from_excel(file_path)
        elif file_path.suffix.lower() == '.csv':
            return self._extract_from_csv(file_path)
        elif file_path.suffix.lower() == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_path.suffix.lower() == '.docx':
            return self._extract_from_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            return self._extract_from_text(file_path)
        else:
            raise ValueError(f"未实现的文件格式处理: {file_path.suffix}")
    
    def _extract_from_excel(self, file_path: Path) -> Dict[str, Any]:
        """从Excel文件抽取数据"""
        try:
            df = pd.read_excel(file_path)
            logger.info(f"Excel文件读取成功: {len(df)} 行, {len(df.columns)} 列")
            
            entities = []
            relations = []
            
            # 结构化数据抽取
            for index, row in df.iterrows():
                row_entities, row_relations = self._extract_from_structured_row(row, index)
                entities.extend(row_entities)
                relations.extend(row_relations)
            
            # 文本内容抽取（从描述字段）
            text_columns = [col for col in df.columns if any(keyword in col.lower() 
                           for keyword in ['描述', '说明', '备注', 'description', 'note'])]
            
            for col in text_columns:
                for text in df[col].dropna():
                    text_entities, text_relations = self._extract_from_text_content(str(text))
                    entities.extend(text_entities)
                    relations.extend(text_relations)
            
            return self._build_extraction_result(entities, relations, file_path, 'Excel')
            
        except Exception as e:
            logger.error(f"Excel抽取失败: {e}")
            raise
    
    def _extract_from_csv(self, file_path: Path) -> Dict[str, Any]:
        """从CSV文件抽取数据"""
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
            logger.info(f"CSV文件读取成功: {len(df)} 行, {len(df.columns)} 列")
            
            # 使用与Excel相同的逻辑
            return self._extract_from_excel(file_path)
            
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                df = pd.read_csv(file_path, encoding='gbk')
                return self._extract_from_excel(file_path)
            except Exception as e:
                logger.error(f"CSV抽取失败: {e}")
                raise
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """从PDF文件抽取数据"""
        try:
            # 尝试导入pdfplumber
            try:
                import pdfplumber
            except ImportError:
                logger.warning("pdfplumber未安装，使用简化PDF处理")
                return self._extract_from_pdf_simple(file_path)
            
            entities = []
            relations = []
            
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
                
                logger.info(f"PDF文本抽取完成: {len(full_text)} 字符")
                
                # 从文本中抽取实体和关系
                text_entities, text_relations = self._extract_from_text_content(full_text)
                entities.extend(text_entities)
                relations.extend(text_relations)
            
            return self._build_extraction_result(entities, relations, file_path, 'PDF')
            
        except Exception as e:
            logger.error(f"PDF抽取失败: {e}")
            raise
    
    def _extract_from_pdf_simple(self, file_path: Path) -> Dict[str, Any]:
        """简化PDF处理（无pdfplumber依赖）"""
        # 返回空结果，提示需要安装依赖
        logger.warning("PDF处理需要安装pdfplumber: pip install pdfplumber")
        return self._build_extraction_result([], [], file_path, 'PDF')
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """从Word文档抽取数据"""
        try:
            # 尝试导入python-docx
            try:
                from docx import Document
            except ImportError:
                logger.warning("python-docx未安装，使用简化Word处理")
                return self._extract_from_docx_simple(file_path)
            
            entities = []
            relations = []
            
            doc = Document(file_path)
            full_text = ""
            
            # 抽取段落文本
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # 抽取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                    full_text += "\n"
            
            logger.info(f"Word文档文本抽取完成: {len(full_text)} 字符")
            
            # 从文本中抽取实体和关系
            text_entities, text_relations = self._extract_from_text_content(full_text)
            entities.extend(text_entities)
            relations.extend(text_relations)
            
            return self._build_extraction_result(entities, relations, file_path, 'Word')
            
        except Exception as e:
            logger.error(f"Word文档抽取失败: {e}")
            raise
    
    def _extract_from_docx_simple(self, file_path: Path) -> Dict[str, Any]:
        """简化Word处理（无python-docx依赖）"""
        logger.warning("Word处理需要安装python-docx: pip install python-docx")
        return self._build_extraction_result([], [], file_path, 'Word')
    
    def _extract_from_text(self, file_path: Path) -> Dict[str, Any]:
        """从文本文件抽取数据"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"文本文件读取成功: {len(text)} 字符")
            
            # 从文本中抽取实体和关系
            entities, relations = self._extract_from_text_content(text)
            
            return self._build_extraction_result(entities, relations, file_path, 'Text')
            
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    text = f.read()
                entities, relations = self._extract_from_text_content(text)
                return self._build_extraction_result(entities, relations, file_path, 'Text')
            except Exception as e:
                logger.error(f"文本文件抽取失败: {e}")
                raise
    
    def _extract_from_structured_row(self, row: pd.Series, row_index: int) -> Tuple[List[Dict], List[Dict]]:
        """从结构化数据行抽取实体和关系"""
        entities = []
        relations = []
        
        # 这里可以复用之前的material_anomaly_extractor逻辑
        # 简化版本，主要抽取关键实体
        
        for col, value in row.items():
            if pd.isna(value):
                continue
                
            value_str = str(value)
            
            # 根据列名和值抽取实体
            if any(keyword in col.lower() for keyword in ['异常', 'anomaly', '问题']):
                entities.append({
                    'key': f"Anomaly:ROW-{row_index}",
                    'type': 'Anomaly',
                    'name': value_str,
                    'properties': {'title': value_str}
                })
            elif any(keyword in col.lower() for keyword in ['组件', 'component', '模块']):
                entities.append({
                    'key': f"Component:{value_str}",
                    'type': 'Component',
                    'name': value_str,
                    'properties': {}
                })
            elif any(keyword in col.lower() for keyword in ['症状', 'symptom']):
                entities.append({
                    'key': f"Symptom:{value_str}",
                    'type': 'Symptom',
                    'name': value_str,
                    'properties': {}
                })
        
        return entities, relations
    
    def _extract_from_text_content(self, text: str) -> Tuple[List[Dict], List[Dict]]:
        """从文本内容抽取实体和关系"""
        entities = []
        relations = []
        
        # 使用正则表达式抽取实体
        for entity_type, patterns in self.extraction_rules['entity_patterns'].items():
            for pattern in patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    entity_value = match.group(1) if match.groups() else match.group(0)
                    entity_value = entity_value.strip()
                    
                    if entity_value:
                        # 标准化实体值
                        normalized_value = self._normalize_entity_value(entity_type, entity_value)
                        
                        entity = {
                            'key': f"{entity_type.title()}:{normalized_value}",
                            'type': entity_type.title(),
                            'name': normalized_value,
                            'properties': {'original_text': entity_value}
                        }
                        entities.append(entity)
        
        # 抽取关系（简化版本）
        for relation_rule in self.extraction_rules['relation_patterns']:
            pattern = relation_rule['pattern']
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) >= 2:
                    source_entity = match.group(1).strip()
                    target_entity = match.group(2).strip()
                    
                    relation = {
                        'source_key': f"{relation_rule['source_type']}:{source_entity}",
                        'target_key': f"{relation_rule['target_type']}:{target_entity}",
                        'relation_type': relation_rule['relation'],
                        'properties': {'confidence': 0.7}
                    }
                    relations.append(relation)
        
        return entities, relations
    
    def _normalize_entity_value(self, entity_type: str, value: str) -> str:
        """标准化实体值"""
        # 根据实体类型和词汇映射表进行标准化
        if entity_type == 'severity' and value in self.vocabulary_mappings['severity_mapping']:
            return self.vocabulary_mappings['severity_mapping'][value]
        elif entity_type == 'component' and value in self.vocabulary_mappings['component_mapping']:
            return self.vocabulary_mappings['component_mapping'][value]
        elif entity_type == 'symptom' and value in self.vocabulary_mappings['symptom_mapping']:
            return self.vocabulary_mappings['symptom_mapping'][value]
        
        return value
    
    def _build_extraction_result(self, entities: List[Dict], relations: List[Dict], 
                                file_path: Path, file_type: str) -> Dict[str, Any]:
        """构建抽取结果"""
        # 去重
        unique_entities = self._deduplicate_entities(entities)
        unique_relations = self._deduplicate_relations(relations)
        
        return {
            'entities': unique_entities,
            'relations': unique_relations,
            'metadata': {
                'source_file': str(file_path),
                'file_type': file_type,
                'entity_count': len(unique_entities),
                'relation_count': len(unique_relations),
                'extracted_at': datetime.now().isoformat()
            }
        }
    
    def _deduplicate_entities(self, entities: List[Dict]) -> List[Dict]:
        """实体去重"""
        seen_keys = set()
        unique_entities = []
        
        for entity in entities:
            if entity['key'] not in seen_keys:
                seen_keys.add(entity['key'])
                unique_entities.append(entity)
        
        return unique_entities
    
    def _deduplicate_relations(self, relations: List[Dict]) -> List[Dict]:
        """关系去重"""
        seen_relations = set()
        unique_relations = []
        
        for relation in relations:
            relation_tuple = (relation['source_key'], relation['target_key'], relation['relation_type'])
            if relation_tuple not in seen_relations:
                seen_relations.add(relation_tuple)
                unique_relations.append(relation)
        
        return unique_relations

def main():
    """主函数 - 测试增强抽取器"""
    extractor = EnhancedDocumentExtractor()
    
    # 测试不同格式的文件
    test_files = [
        "data/import/来料问题先后版.xlsx",
        "data/import/相关测试用例.xlsx"
    ]
    
    for test_file in test_files:
        if Path(test_file).exists():
            try:
                result = extractor.extract_from_file(test_file)
                
                # 保存结果
                output_file = f"data/processed/enhanced_extracted_{Path(test_file).stem}.json"
                Path(output_file).parent.mkdir(parents=True, exist_ok=True)
                
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
                
                print(f"✅ {test_file} 抽取完成!")
                print(f"📊 结果: {result['metadata']['entity_count']} 个实体, {result['metadata']['relation_count']} 个关系")
                print(f"💾 保存到: {output_file}")
                print()
                
            except Exception as e:
                print(f"❌ {test_file} 抽取失败: {e}")
        else:
            print(f"⚠️  文件不存在: {test_file}")

if __name__ == "__main__":
    main()

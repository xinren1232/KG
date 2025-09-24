#!/usr/bin/env python3
"""
文件信息抽取器
支持多种格式文件的内容提取和结构化处理
"""

import os
import re
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, asdict
import pandas as pd
import pdfplumber
from docx import Document
import openpyxl
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)

@dataclass
class ExtractedEntity:
    """抽取的实体"""
    id: str
    key: str
    name: str
    type: str  # product, component, test_case, anomaly, etc.
    properties: Dict[str, Any]
    source_file: str
    source_location: str  # 在文件中的位置信息

@dataclass
class ExtractedRelation:
    """抽取的关系"""
    source_entity: str
    target_entity: str
    relation_type: str
    properties: Dict[str, Any]
    source_file: str
    confidence: float = 1.0

@dataclass
class ExtractionResult:
    """抽取结果"""
    file_path: str
    file_type: str
    entities: List[ExtractedEntity]
    relations: List[ExtractedRelation]
    metadata: Dict[str, Any]
    errors: List[str]

class FileExtractor:
    """文件抽取器基类"""

    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.csv', '.pdf', '.docx', '.txt']
        self.entity_patterns = self._load_entity_patterns()
        self.relation_patterns = self._load_relation_patterns()
        # 标准标签映射
        self.type_to_label = {
            'product': 'Product',
            'build': 'Build',
            'component': 'Component',
            'test_case': 'TestCase',
            'test_step': 'TestStep',
            'test_result': 'TestResult',
            'anomaly': 'Anomaly',
            'symptom': 'Symptom',
            'root_cause': 'RootCause',
            'countermeasure': 'Countermeasure',
            'owner': 'Owner',
            'supplier': 'Supplier',
            'doc': 'Doc',
        }
        # 加载词表
        self._load_vocabs()
    def _load_vocabs(self) -> None:
        """加载标准词表：组件别名->标准名，症状/根因标准集合"""
        self.comp_alias2std: Dict[str, str] = {}
        self.symptom_std: List[str] = []
        self.cause_std: List[str] = []
        try:
            comp_path = Path('data/vocab/components.csv')
            if comp_path.exists():
                dfc = pd.read_csv(comp_path)
                for _, r in dfc.iterrows():
                    name = str(r.get('name', '')).strip()
                    alias = str(r.get('alias', '')).strip()
                    if name:
                        self.comp_alias2std[name.lower()] = name
                    if alias:
                        self.comp_alias2std[alias.lower()] = name or alias
        except Exception as e:
            logger.warning(f"加载组件词表失败: {e}")
        try:
            sym_path = Path('data/vocab/symptoms.csv')
            if sym_path.exists():
                dfs = pd.read_csv(sym_path)
                self.symptom_std = [str(x).strip() for x in dfs['name'].dropna().tolist()]
        except Exception as e:
            logger.warning(f"加载症状词表失败: {e}")
        try:
            cause_path = Path('data/vocab/causes.csv')
            if cause_path.exists():
                dfcau = pd.read_csv(cause_path)
                self.cause_std = [str(x).strip() for x in dfcau['name'].dropna().tolist()]
        except Exception as e:
            logger.warning(f"加载根因词表失败: {e}")

    def _standardize_name(self, entity_type: str, name: str) -> str:
        """根据词表标准化名称（中英别名兜底）"""
        if not name:
            return name
        n = name.strip()
        t = (entity_type or '').lower()
        if t == 'component' and self.comp_alias2std:
            return self.comp_alias2std.get(n.lower(), n)
        if t == 'symptom' and self.symptom_std:
            for std in self.symptom_std:
                if n.lower() == std.lower():
                    return std
            return n
        if t == 'root_cause' and self.cause_std:
            for std in self.cause_std:
                if n.lower() == std.lower():
                    return std
            return n
        return n

    def _make_key(self, label: str, name: str, extra: Optional[Dict[str, Any]] = None) -> str:
        """生成统一业务主键 key"""
        extra = extra or {}
        if label in ('Product','Component','Owner','Supplier','Doc','Symptom','RootCause','Countermeasure'):
            return f"{label}:{name}"
        if label == 'Build':
            return f"Build:{extra.get('version', name)}"
        if label == 'TestCase':
            return f"TestCase:{name}"
        if label == 'TestStep':
            cid = extra.get('case_id', name)
            idx = extra.get('index', '')
            return f"TestStep:{cid}-{idx}".rstrip('-')
        if label == 'TestResult':
            suffix = extra.get('build') or extra.get('version') or ''
            return f"TestResult:{name}-{suffix}".rstrip('-')
        if label == 'Anomaly':
            return f"Anomaly:{extra.get('code', name)}"
        return f"{label}:{name}"


    def _load_entity_patterns(self) -> Dict[str, List[str]]:
        """加载实体识别模式"""
        return {
            'product': [
                r'iPhone\s*\d+',
                r'iPad\s*\w+',
                r'MacBook\s*\w+',
                r'产品[：:]\s*([^\n\r,，]+)',
                r'机型[：:]\s*([^\n\r,，]+)'
            ],
            'component': [
                r'摄像头|相机|Camera',
                r'屏幕|显示屏|Display',
                r'电池|Battery',
                r'处理器|CPU|芯片',
                r'内存|Memory|RAM',
                r'存储|Storage|闪存',
                r'扬声器|Speaker',
                r'麦克风|Microphone',
                r'传感器|Sensor',
                r'组件[：:]\s*([^\n\r,，]+)'
            ],
            'test_case': [
                r'TC[-_]\d+',
                r'测试用例[-_]\d+',
                r'Test\s*Case\s*\d+',
                r'用例编号[：:]\s*([^\n\r,，]+)'
            ],
            'anomaly': [
                r'异常[-_]\d+',
                r'Bug[-_]\d+',
                r'缺陷[-_]\d+',
                r'问题[-_]\d+',
                r'故障[：:]\s*([^\n\r,，]+)'
            ]
        }

    def _load_relation_patterns(self) -> Dict[str, List[str]]:
        """加载关系识别模式"""
        return {
            'contains': [
                r'包含|包括|含有',
                r'consists of|contains|includes'
            ],
            'tests': [
                r'测试|检测|验证',
                r'test|verify|validate'
            ],
            'affects': [
                r'影响|导致|造成',
                r'affects|causes|leads to'
            ],
            'belongs_to': [
                r'属于|归属于',
                r'belongs to|part of'
            ]
        }

    def extract_file(self, file_path: str) -> ExtractionResult:
        """抽取文件内容"""
        file_path = Path(file_path)

        if not file_path.exists():
            return ExtractionResult(
                file_path=str(file_path),
                file_type='unknown',
                entities=[],
                relations=[],
                metadata={},
                errors=[f"文件不存在: {file_path}"]
            )

        file_ext = file_path.suffix.lower()

        try:
            if file_ext in ['.xlsx', '.xls']:
                return self._extract_excel(file_path)
            elif file_ext == '.csv':
                return self._extract_csv(file_path)
            elif file_ext == '.pdf':
                return self._extract_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_docx(file_path)
            elif file_ext == '.txt':
                return self._extract_txt(file_path)
            else:
                return ExtractionResult(
                    file_path=str(file_path),
                    file_type=file_ext,
                    entities=[],
                    relations=[],
                    metadata={},
                    errors=[f"不支持的文件格式: {file_ext}"]
                )

        except Exception as e:
            logger.error(f"文件抽取失败 {file_path}: {e}")
            return ExtractionResult(
                file_path=str(file_path),
                file_type=file_ext,
                entities=[],
                relations=[],
                metadata={},
                errors=[f"抽取失败: {str(e)}"]
            )

    def _extract_excel(self, file_path: Path) -> ExtractionResult:
        """抽取Excel文件"""
        entities = []
        relations = []
        errors = []

        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}

            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    sheets_data[sheet_name] = df

                    # 从表格中抽取实体
                    sheet_entities, sheet_relations = self._extract_from_dataframe(
                        df, str(file_path), f"Sheet:{sheet_name}"
                    )
                    entities.extend(sheet_entities)
                    relations.extend(sheet_relations)

                except Exception as e:
                    errors.append(f"读取工作表 {sheet_name} 失败: {str(e)}")

            metadata = {
                'sheets': list(sheets_data.keys()),
                'total_rows': sum(len(df) for df in sheets_data.values()),
                'file_size': file_path.stat().st_size
            }

        except Exception as e:
            errors.append(f"Excel文件处理失败: {str(e)}")
            metadata = {}

        return ExtractionResult(
            file_path=str(file_path),
            file_type='excel',
            entities=entities,
            relations=relations,
            metadata=metadata,
            errors=errors
        )

    def _extract_csv(self, file_path: Path) -> ExtractionResult:
        """抽取CSV文件"""
        entities = []
        relations = []
        errors = []

        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']
            df = None

            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                errors.append("无法解码CSV文件，尝试了多种编码格式")
                return ExtractionResult(
                    file_path=str(file_path),
                    file_type='csv',
                    entities=[],
                    relations=[],
                    metadata={},
                    errors=errors
                )

            # 从DataFrame抽取实体和关系
            entities, relations = self._extract_from_dataframe(
                df, str(file_path), "CSV"
            )

            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'file_size': file_path.stat().st_size
            }

        except Exception as e:
            errors.append(f"CSV文件处理失败: {str(e)}")
            metadata = {}

        return ExtractionResult(
            file_path=str(file_path),
            file_type='csv',
            entities=entities,
            relations=relations,
            metadata=metadata,
            errors=errors
        )

    def _extract_pdf(self, file_path: Path) -> ExtractionResult:
        """抽取PDF文件"""
        entities = []
        relations = []
        errors = []
        text_content = ""

        try:
            with pdfplumber.open(file_path) as pdf:
                # 提取文本内容
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"

                        # 从页面文本中抽取实体
                        page_entities, page_relations = self._extract_from_text(
                            page_text, str(file_path), f"Page:{page_num+1}"
                        )
                        entities.extend(page_entities)
                        relations.extend(page_relations)

                metadata = {
                    'pages': len(pdf.pages),
                    'text_length': len(text_content),
                    'file_size': file_path.stat().st_size
                }

        except Exception as e:
            errors.append(f"PDF文件处理失败: {str(e)}")
            metadata = {}

        return ExtractionResult(
            file_path=str(file_path),
            file_type='pdf',
            entities=entities,
            relations=relations,
            metadata=metadata,
            errors=errors
        )

    def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """抽取Word文档"""
        entities = []
        relations = []
        errors = []
        text_content = ""

        try:
            doc = Document(file_path)

            # 提取段落文本
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"

                    # 从段落中抽取实体
                    para_entities, para_relations = self._extract_from_text(
                        paragraph.text, str(file_path), f"Paragraph:{para_num+1}"
                    )
                    entities.extend(para_entities)
                    relations.extend(para_relations)

            # 提取表格
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    # 转换为DataFrame进行处理
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    table_entities, table_relations = self._extract_from_dataframe(
                        df, str(file_path), f"Table:{table_num+1}"
                    )
                    entities.extend(table_entities)
                    relations.extend(table_relations)

            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'text_length': len(text_content),
                'file_size': file_path.stat().st_size
            }

        except Exception as e:
            errors.append(f"Word文档处理失败: {str(e)}")
            metadata = {}

        return ExtractionResult(
            file_path=str(file_path),
            file_type='docx',
            entities=entities,
            relations=relations,
            metadata=metadata,
            errors=errors
        )

    def _extract_txt(self, file_path: Path) -> ExtractionResult:
        """抽取文本文件"""
        entities = []
        relations = []
        errors = []

        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']
            text_content = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text_content is None:
                errors.append("无法解码文本文件，尝试了多种编码格式")
                return ExtractionResult(
                    file_path=str(file_path),
                    file_type='txt',
                    entities=[],
                    relations=[],
                    metadata={},
                    errors=errors
                )

            # 从文本中抽取实体和关系
            entities, relations = self._extract_from_text(
                text_content, str(file_path), "Text"
            )

            metadata = {
                'text_length': len(text_content),
                'lines': len(text_content.split('\n')),
                'file_size': file_path.stat().st_size
            }

        except Exception as e:
            errors.append(f"文本文件处理失败: {str(e)}")
            metadata = {}

        return ExtractionResult(
            file_path=str(file_path),
            file_type='txt',
            entities=entities,
            relations=relations,
            metadata=metadata,
            errors=errors
        )

    def _extract_from_dataframe(self, df: pd.DataFrame, source_file: str, location: str) -> Tuple[List[ExtractedEntity], List[ExtractedRelation]]:
        """从DataFrame中抽取实体和关系"""
        entities = []
        relations = []

        # 检查列名，识别可能的实体类型
        column_mapping = self._map_columns_to_entity_types(df.columns)

        for idx, row in df.iterrows():
            row_entities = []

            # 从每一行抽取实体
            for col, entity_type in column_mapping.items():
                if col in df.columns and pd.notna(row[col]):
                    value = str(row[col]).strip()
                    if value:
                        label = self.type_to_label.get(entity_type, entity_type.title())
                        std_name = self._standardize_name(entity_type, value)
                        entity_id = f"{entity_type}_{hash(std_name) % 10000:04d}"
                        key = self._make_key(label, std_name, {'row': idx})
                        entity = ExtractedEntity(
                            id=entity_id,
                            key=key,
                            name=std_name,
                            type=entity_type,
                            properties={'original_value': value, 'label': label, 'row': idx},
                            source_file=source_file,
                            source_location=f"{location}:Row{idx+1}:Col{col}"
                        )
                        entities.append(entity)
                        row_entities.append(entity)

            # 行内基于领域规则建立关系
            if row_entities:
                # 按类型分类
                by_type: Dict[str, List[ExtractedEntity]] = {}
                for e in row_entities:
                    by_type.setdefault(e.type, []).append(e)
                prod = by_type.get('product', [None])[0]
                # Product -> Component
                if prod:
                    for comp in by_type.get('component', []):
                        relations.append(ExtractedRelation(
                            source_entity=prod.id,
                            target_entity=comp.id,
                            relation_type='INCLUDES',
                            properties={'context': 'same_row'},
                            source_file=source_file,
                            confidence=0.9
                        ))
                    # BELONGS_TO for lower-level entities
                    for t in ['component', 'test_case', 'anomaly']:
                        for ent in by_type.get(t, []):
                            relations.append(ExtractedRelation(
                                source_entity=ent.id,
                                target_entity=prod.id,
                                relation_type='BELONGS_TO',
                                properties={'context': 'same_row'},
                                source_file=source_file,
                                confidence=0.8
                            ))
                # Anomaly relations
                if by_type.get('anomaly'):
                    a = by_type['anomaly'][0]
                    for s in by_type.get('symptom', []):
                        relations.append(ExtractedRelation(
                            source_entity=a.id,
                            target_entity=s.id,
                            relation_type='HAS_SYMPTOM',
                            properties={'context': 'same_row'},
                            source_file=source_file,
                            confidence=0.9
                        ))
                    for cm in by_type.get('countermeasure', []):
                        relations.append(ExtractedRelation(
                            source_entity=a.id,
                            target_entity=cm.id,
                            relation_type='RESOLVED_BY',
                            properties={'context': 'same_row'},
                            source_file=source_file,
                            confidence=0.8
                        ))
                    for rc in by_type.get('root_cause', []):
                        relations.append(ExtractedRelation(
                            source_entity=rc.id,
                            target_entity=a.id,
                            relation_type='CAUSES',
                            properties={'context': 'same_row'},
                            source_file=source_file,
                            confidence=0.85
                        ))

        return entities, relations

    def _extract_from_text(self, text: str, source_file: str, location: str) -> Tuple[List[ExtractedEntity], List[ExtractedRelation]]:
        """从文本中抽取实体和关系"""
        entities = []
        relations = []

        # 使用正则表达式抽取实体
        for entity_type, patterns in self.entity_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    entity_name = match.group(1) if match.groups() else match.group(0)
                    entity_name = entity_name.strip()

                    if entity_name:
                        entity_id = f"{entity_type}_{hash(entity_name) % 10000:04d}"
                        label = self.type_to_label.get(entity_type, entity_type.title())
                        std_name = self._standardize_name(entity_type, entity_name)
                        key = self._make_key(label, std_name, {})
                        entity = ExtractedEntity(
                            id=entity_id,
                            key=key,
                            name=std_name,
                            type=entity_type,
                            properties={'pattern': pattern, 'match_text': match.group(0), 'label': label},
                            source_file=source_file,
                            source_location=f"{location}:Pos{match.start()}-{match.end()}"
                        )
                        entities.append(entity)

        # 简单的关系抽取（基于共现）

        for i, entity1 in enumerate(entities):
            for entity2 in entities[i+1:]:
                # 检查两个实体是否在相近的位置出现
                if self._entities_are_related(entity1.name, entity2.name, text):
                    relation = ExtractedRelation(
                        source_entity=entity1.id,
                        target_entity=entity2.id,
                        relation_type='co_occurs',
                        properties={'context': 'text_proximity'},
                        source_file=source_file,
                        confidence=0.6
                    )
                    relations.append(relation)

        return entities, relations

    def _map_columns_to_entity_types(self, columns: List[str]) -> Dict[str, str]:
        """将列名映射到实体类型"""
        mapping = {}

        for col in columns:
            col_lower = col.lower()
            if any(keyword in col_lower for keyword in ['产品', 'product', '机型', 'model']):
                mapping[col] = 'product'
            elif any(keyword in col_lower for keyword in ['组件', 'component', '模块', 'module']):
                mapping[col] = 'component'
            elif any(keyword in col_lower for keyword in ['测试', 'test', '用例', 'case']):
                mapping[col] = 'test_case'
            elif any(keyword in col_lower for keyword in ['异常', 'anomaly', '缺陷', 'bug', '问题', 'issue']):
                mapping[col] = 'anomaly'
            elif any(keyword in col_lower for keyword in ['症状', 'symptom', '现象']):
                mapping[col] = 'symptom'
            elif any(keyword in col_lower for keyword in ['原因', 'cause', '根因']):
                mapping[col] = 'root_cause'
            elif any(keyword in col_lower for keyword in ['对策', 'solution', '解决', 'fix']):
                mapping[col] = 'countermeasure'

        return mapping

    def _entities_are_related(self, entity1: str, entity2: str, text: str) -> bool:
        """判断两个实体是否相关（基于文本距离）"""
        # 简单的距离判断：如果两个实体在同一句话中出现，认为相关
        sentences = re.split(r'[。！？\n]', text)

        for sentence in sentences:
            if entity1 in sentence and entity2 in sentence:
                return True

        return False

# 使用示例
if __name__ == "__main__":
    extractor = FileExtractor()

    # 测试文件抽取
    test_files = [
        "../../data/raw/测试用例样本数据.xlsx",
        "../../data/raw/异常数据样本.xlsx"
    ]

    for file_path in test_files:
        if os.path.exists(file_path):
            result = extractor.extract_file(file_path)
            print(f"\n文件: {result.file_path}")
            print(f"类型: {result.file_type}")
            print(f"实体数量: {len(result.entities)}")
            print(f"关系数量: {len(result.relations)}")
            print(f"错误: {result.errors}")

            # 显示前几个实体
            for entity in result.entities[:3]:
                print(f"  实体: {entity.name} ({entity.type})")
        else:
            print(f"文件不存在: {file_path}")
